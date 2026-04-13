"""
BountyLens MCP Server v2
========================
MCP server with comprehensive checklist engine integration.

New in v2:
- OWASP API Top 10 mapped test cases
- HackerOne/Bugcrowd/Hacktify bounty report patterns
- BOLA/BFLA specific deep-dive tests
- Custom checklist support
- Auto-selection based on endpoint parameter analysis
- Three views: per-endpoint, per-parameter, per-vulnerability-class
- Toggle test cases on/off
- Per-parameter test tracking

Usage:
    pip install "mcp[cli]" httpx python-docx reportlab
    python bountylens_mcp.py
"""

import json
import threading
import uuid
from datetime import datetime
from http.server import HTTPServer, BaseHTTPRequestHandler
from typing import Any

from mcp.server.fastmcp import FastMCP

from checklist_engine import (
    MASTER_CHECKLIST,
    OWASP_CATEGORIES,
    auto_select_tests,
    get_per_parameter_view,
    get_per_vuln_class_view,
    test_case_stats,
    custom_checklists,
)

# ──────────────────────────────────────────────
# In-memory data store
# ──────────────────────────────────────────────

endpoints_db: dict[str, dict] = {}

# ──────────────────────────────────────────────
# HTTP Bridge for Burp Extension
# ──────────────────────────────────────────────

BRIDGE_PORT = 8888


class BurpBridgeHandler(BaseHTTPRequestHandler):
    def do_POST(self):
        if self.path == "/endpoints":
            content_length = int(self.headers.get("Content-Length", 0))
            body = self.rfile.read(content_length)
            try:
                data = json.loads(body)
                endpoint_id = str(uuid.uuid4())[:8]
                endpoints_db[endpoint_id] = _create_endpoint(endpoint_id, data)
                self._respond(200, {"status": "ok", "id": endpoint_id})
            except json.JSONDecodeError:
                self._respond(400, {"error": "Invalid JSON"})

        elif self.path == "/endpoints/bulk":
            content_length = int(self.headers.get("Content-Length", 0))
            body = self.rfile.read(content_length)
            try:
                data_list = json.loads(body)
                ids = []
                for data in data_list:
                    endpoint_id = str(uuid.uuid4())[:8]
                    endpoints_db[endpoint_id] = _create_endpoint(endpoint_id, data)
                    ids.append(endpoint_id)
                self._respond(200, {"status": "ok", "ids": ids, "count": len(ids)})
            except json.JSONDecodeError:
                self._respond(400, {"error": "Invalid JSON"})

        elif self.path == "/clear":
            endpoints_db.clear()
            self._respond(200, {"status": "cleared"})
        else:
            self._respond(404, {"error": "Not found"})

    def do_GET(self):
        if self.path == "/health":
            self._respond(200, {"status": "healthy", "endpoints_count": len(endpoints_db)})
        elif self.path == "/endpoints":
            self._respond(200, list(endpoints_db.values()))
        else:
            self._respond(404, {"error": "Not found"})

    def _respond(self, code: int, data: Any):
        self.send_response(code)
        self.send_header("Content-Type", "application/json")
        self.send_header("Access-Control-Allow-Origin", "*")
        self.end_headers()
        self.wfile.write(json.dumps(data, default=str).encode())

    def log_message(self, format, *args):
        pass


def _create_endpoint(endpoint_id: str, data: dict) -> dict:
    ep = {
        "id": endpoint_id,
        "url": data.get("url", ""),
        "method": data.get("method", "GET"),
        "headers": data.get("headers", {}),
        "query_params": data.get("query_params", {}),
        "body_params": data.get("body_params", {}),
        "content_type": data.get("content_type", ""),
        "discovered_at": datetime.now().isoformat(),
        "test_cases": [],
        "business_context": "",
        "risk_level": "unassessed",
    }
    ep["test_cases"] = auto_select_tests(ep)
    return ep


def start_bridge_server():
    server = HTTPServer(("127.0.0.1", BRIDGE_PORT), BurpBridgeHandler)
    thread = threading.Thread(target=server.serve_forever, daemon=True)
    thread.start()
    return server


# ──────────────────────────────────────────────
# MCP Server
# ──────────────────────────────────────────────

mcp = FastMCP(
    "BountyLens",
    description=(
        "AI-powered API Security Testing Platform — crawl, analyze, test, track, report. "
        "Features OWASP API Top 10 checklists, bounty report patterns, BOLA/BFLA deep-dive, "
        "three views (per-endpoint, per-parameter, per-vuln-class), and smart auto-selection."
    ),
)


@mcp.tool()
def list_endpoints(status_filter: str = "all") -> str:
    """
    List all discovered API endpoints.
    Args:
        status_filter: "all", "tested", "untested", "has_failures"
    """
    if not endpoints_db:
        return "No endpoints discovered yet."

    results = []
    for ep in endpoints_db.values():
        enabled_tests = [tc for tc in ep["test_cases"] if tc.get("enabled")]
        total = len(enabled_tests)
        done = sum(1 for tc in enabled_tests if tc["status"] in ("pass", "fail", "na", "skipped"))
        failed = sum(1 for tc in enabled_tests if tc["status"] == "fail")

        if status_filter == "tested" and done < total:
            continue
        if status_filter == "untested" and done > 0:
            continue
        if status_filter == "has_failures" and failed == 0:
            continue

        results.append(
            f"[{ep['id']}] {ep['method']} {ep['url']}\n"
            f"    Params: {len(ep['query_params'])}q / {len(ep['body_params'])}b | "
            f"Tests: {done}/{total} | Failures: {failed} | Risk: {ep['risk_level']}"
        )

    return f"Found {len(results)} endpoints:\n\n" + "\n\n".join(results)


@mcp.tool()
def get_endpoint_details(endpoint_id: str) -> str:
    """Get full details of a specific endpoint."""
    ep = endpoints_db.get(endpoint_id)
    if not ep:
        return f"Endpoint '{endpoint_id}' not found."

    stats = test_case_stats(ep["test_cases"], enabled_only=True)
    return (
        f"Endpoint: {ep['method']} {ep['url']}\n"
        f"ID: {ep['id']} | Content-Type: {ep['content_type'] or 'N/A'}\n"
        f"Risk: {ep['risk_level']} | Context: {ep['business_context'] or 'Not set'}\n\n"
        f"Headers: {json.dumps(ep['headers'], indent=2)}\n"
        f"Query: {json.dumps(ep['query_params'], indent=2)}\n"
        f"Body: {json.dumps(ep['body_params'], indent=2)}\n\n"
        f"Tests: {stats['enabled']} enabled | Pass: {stats['passed']} | Fail: {stats['failed']} | "
        f"Pending: {stats['pending']} | Coverage: {stats['coverage_pct']:.1f}%"
    )


@mcp.tool()
def generate_checklist(endpoint_id: str) -> str:
    """
    Generate/regenerate the full security checklist for an endpoint.
    Auto-selects relevant tests based on parameters, headers, method, URL patterns.
    """
    ep = endpoints_db.get(endpoint_id)
    if not ep:
        return f"Endpoint '{endpoint_id}' not found."

    ep["test_cases"] = auto_select_tests(ep)
    total = len(ep["test_cases"])
    auto = sum(1 for tc in ep["test_cases"] if tc["auto_selected"])

    by_cat = {}
    for tc in ep["test_cases"]:
        cat = tc.get("category", "Other")
        if cat not in by_cat:
            by_cat[cat] = []
        by_cat[cat].append(tc)

    output = f"Checklist for {ep['method']} {ep['url']}: {total} tests ({auto} auto-enabled)\n"

    for cat, tests in by_cat.items():
        enabled = sum(1 for t in tests if t["enabled"])
        output += f"\n{'═' * 50}\n  {cat} ({enabled}/{len(tests)} enabled)\n{'═' * 50}\n"
        for tc in tests:
            icon = "✅" if tc["enabled"] else "⬜"
            match = f" ← {', '.join(tc['match_reasons'])}" if tc["match_reasons"] else ""
            output += f"  {icon} [{tc['id']}] {tc['title']}{match}\n"
            output += f"     {tc['severity'].upper()} | {tc['source']}\n"

    return output


@mcp.tool()
def view_checklist(endpoint_id: str, view: str = "endpoint") -> str:
    """
    View checklist in three organizations.
    Args:
        endpoint_id: The endpoint ID
        view: "endpoint" | "parameter" | "vulnerability"
    """
    ep = endpoints_db.get(endpoint_id)
    if not ep:
        return f"Endpoint '{endpoint_id}' not found."
    if not ep["test_cases"]:
        return f"No test cases. Run generate_checklist('{endpoint_id}') first."

    output = f"Checklist — {ep['method']} {ep['url']} — View: {view.upper()}\n\n"
    tc_lookup = {tc["id"]: tc for tc in ep["test_cases"]}

    if view == "parameter":
        param_view = get_per_parameter_view(ep["test_cases"], ep)
        for key, info in param_view.items():
            if not info["tests"]:
                continue
            output += f"\n{'─' * 50}\n  {info['type'].upper()}: {info['name']} ({len(info['tests'])} tests)\n{'─' * 50}\n"
            for tc_id in info["tests"]:
                tc = tc_lookup.get(tc_id)
                if tc:
                    icon = _status_icon(tc["status"]) if tc["enabled"] else "⬜"
                    output += f"  {icon} [{tc['id']}] {tc['title']} ({tc['status']})\n"

    elif view == "vulnerability":
        vuln_view = get_per_vuln_class_view(ep["test_cases"])
        for cat, info in vuln_view.items():
            owasp_label = f" [{info['owasp']}]" if info['owasp'] else ""
            output += f"\n{'─' * 50}\n  {cat}{owasp_label} ({len(info['tests'])} tests)\n{'─' * 50}\n"
            for tc_id in info["tests"]:
                tc = tc_lookup.get(tc_id)
                if tc:
                    icon = _status_icon(tc["status"]) if tc["enabled"] else "⬜"
                    output += f"  {icon} [{tc['severity'].upper()}] [{tc['id']}] {tc['title']} ({tc['status']})\n"

    else:
        by_cat = {}
        for tc in ep["test_cases"]:
            cat = tc.get("category", "Other")
            if cat not in by_cat:
                by_cat[cat] = []
            by_cat[cat].append(tc)
        for cat, tests in by_cat.items():
            enabled = [t for t in tests if t["enabled"]]
            done = sum(1 for t in enabled if t["status"] in ("pass", "fail", "na", "skipped"))
            output += f"\n{'─' * 50}\n  {cat} — {done}/{len(enabled)} done\n{'─' * 50}\n"
            for tc in tests:
                icon = _status_icon(tc["status"]) if tc["enabled"] else "⬜"
                output += f"  {icon} [{tc['id']}] {tc['title']} | {tc['status']}\n"

    stats = test_case_stats(ep["test_cases"], enabled_only=True)
    output += f"\nSUMMARY: {stats['done']}/{stats['total']} | Pass:{stats['passed']} Fail:{stats['failed']} Pending:{stats['pending']} | Coverage:{stats['coverage_pct']:.1f}%"
    return output


@mcp.tool()
def toggle_test_case(endpoint_id: str, test_case_id: str, enabled: bool) -> str:
    """Enable or disable a specific test case."""
    ep = endpoints_db.get(endpoint_id)
    if not ep:
        return f"Endpoint '{endpoint_id}' not found."
    for tc in ep["test_cases"]:
        if tc["id"] == test_case_id:
            tc["enabled"] = enabled
            return f"{'ENABLED' if enabled else 'DISABLED'}: [{tc['id']}] {tc['title']}"
    return f"Test case '{test_case_id}' not found."


@mcp.tool()
def bulk_toggle_tests(endpoint_id: str, category: str = "", owasp: str = "", enabled: bool = True) -> str:
    """Enable/disable multiple tests by category or OWASP mapping."""
    ep = endpoints_db.get(endpoint_id)
    if not ep:
        return f"Endpoint '{endpoint_id}' not found."
    count = 0
    for tc in ep["test_cases"]:
        match = (not category and not owasp)
        if category and tc.get("category", "").lower() == category.lower():
            match = True
        if owasp and tc.get("owasp", "").upper() == owasp.upper():
            match = True
        if match:
            tc["enabled"] = enabled
            count += 1
    return f"{'ENABLED' if enabled else 'DISABLED'} {count} tests"


@mcp.tool()
def set_test_result(endpoint_id: str, test_case_id: str, status: str, evidence: str = "", notes: str = "") -> str:
    """
    Update test case result.
    Args:
        status: "pass", "fail", "na", "skipped"
    """
    ep = endpoints_db.get(endpoint_id)
    if not ep:
        return f"Endpoint '{endpoint_id}' not found."
    for tc in ep["test_cases"]:
        if tc["id"] == test_case_id:
            tc["status"] = status
            tc["evidence"] = evidence
            tc["notes"] = notes
            tc["tested_at"] = datetime.now().isoformat()
            return f"{_status_icon(status)} [{tc['id']}] {tc['title']} → {status.upper()}"
    return f"Test case '{test_case_id}' not found."


@mcp.tool()
def bulk_set_results(endpoint_id: str, results: str) -> str:
    """
    Set multiple test results at once.
    Args:
        results: JSON array e.g. [{"id":"BOLA-001","status":"pass"},{"id":"AUTH-001","status":"fail","evidence":"..."}]
    """
    ep = endpoints_db.get(endpoint_id)
    if not ep:
        return f"Endpoint '{endpoint_id}' not found."
    try:
        result_list = json.loads(results)
    except json.JSONDecodeError:
        return "Invalid JSON."
    tc_lookup = {tc["id"]: tc for tc in ep["test_cases"]}
    updated = 0
    output = []
    for r in result_list:
        tc = tc_lookup.get(r.get("id"))
        if tc:
            tc["status"] = r.get("status", "pending")
            tc["evidence"] = r.get("evidence", "")
            tc["notes"] = r.get("notes", "")
            tc["tested_at"] = datetime.now().isoformat()
            updated += 1
            output.append(f"  {_status_icon(tc['status'])} [{tc['id']}] → {tc['status'].upper()}")
    return f"Updated {updated} tests:\n" + "\n".join(output)


@mcp.tool()
def get_test_case_detail(endpoint_id: str, test_case_id: str) -> str:
    """Get full details of a specific test case including payloads."""
    ep = endpoints_db.get(endpoint_id)
    if not ep:
        return f"Endpoint '{endpoint_id}' not found."
    for tc in ep["test_cases"]:
        if tc["id"] == test_case_id:
            return json.dumps({
                "id": tc["id"], "title": tc["title"], "description": tc["description"],
                "category": tc.get("category"), "owasp": tc.get("owasp"),
                "owasp_name": OWASP_CATEGORIES.get(tc.get("owasp", ""), {}).get("name", ""),
                "severity": tc["severity"], "source": tc["source"],
                "payloads": tc.get("payloads", []), "enabled": tc["enabled"],
                "auto_selected": tc.get("auto_selected"), "match_reasons": tc.get("match_reasons", []),
                "target_params": tc.get("target_params", []),
                "status": tc["status"], "evidence": tc["evidence"],
                "notes": tc.get("notes", ""), "tested_at": tc.get("tested_at"),
            }, indent=2)
    return f"Test case '{test_case_id}' not found."


@mcp.tool()
def add_custom_test(title: str, description: str, category: str = "Custom",
                    severity: str = "medium", owasp: str = "",
                    param_triggers: str = "", url_triggers: str = "", payloads: str = "") -> str:
    """
    Add a custom test case to the global checklist.
    Args:
        param_triggers: Comma-separated param names that trigger this (e.g. "price,amount")
        url_triggers: Comma-separated URL patterns (e.g. "checkout,payment")
        payloads: Comma-separated payloads
    """
    tc_id = f"CUSTOM-{len(custom_checklists) + 1:03d}"
    triggers = {}
    if param_triggers:
        triggers["param_names"] = [p.strip() for p in param_triggers.split(",")]
    if url_triggers:
        triggers["url_patterns"] = [u.strip() for u in url_triggers.split(",")]
    if not triggers:
        triggers["always"] = True

    custom_checklists.append({
        "id": tc_id, "title": title, "description": description,
        "owasp": owasp, "source": "Custom", "severity": severity,
        "category": category, "triggers": triggers,
        "payloads": [p.strip() for p in payloads.split(",")] if payloads else [],
    })
    return f"Added [{tc_id}] '{title}'"


@mcp.tool()
def add_business_context(endpoint_id: str, context: str, risk_level: str = "unassessed") -> str:
    """Add business context and risk for an endpoint."""
    ep = endpoints_db.get(endpoint_id)
    if not ep:
        return f"Endpoint '{endpoint_id}' not found."
    ep["business_context"] = context
    ep["risk_level"] = risk_level
    return f"Context set for {ep['method']} {ep['url']} (Risk: {risk_level})"


@mcp.tool()
def get_coverage_dashboard() -> str:
    """Full coverage dashboard with OWASP breakdown and blind spots."""
    if not endpoints_db:
        return "No endpoints."

    total_ep = len(endpoints_db)
    fully = partially = untested_ep = 0
    total_tests = total_enabled = passed = failed = na_c = skipped_c = pending = 0
    blind_spots = []
    owasp_cov = {cat: {"total": 0, "done": 0, "failed": 0} for cat in OWASP_CATEGORIES}

    for ep in endpoints_db.values():
        enabled = [tc for tc in ep["test_cases"] if tc.get("enabled")]
        total_tests += len(ep["test_cases"])
        total_enabled += len(enabled)
        done = 0
        for tc in enabled:
            oc = tc.get("owasp", "")
            if oc in owasp_cov:
                owasp_cov[oc]["total"] += 1
            if tc["status"] == "pass":
                passed += 1; done += 1
                if oc in owasp_cov: owasp_cov[oc]["done"] += 1
            elif tc["status"] == "fail":
                failed += 1; done += 1
                if oc in owasp_cov: owasp_cov[oc]["done"] += 1; owasp_cov[oc]["failed"] += 1
            elif tc["status"] in ("na", "skipped"):
                if tc["status"] == "na": na_c += 1
                else: skipped_c += 1
                done += 1
                if oc in owasp_cov: owasp_cov[oc]["done"] += 1
            else:
                pending += 1

        if not enabled:
            blind_spots.append(f"  ⚠ [{ep['id']}] {ep['method']} {ep['url']} — NO TESTS")
        elif done == len(enabled): fully += 1
        elif done > 0: partially += 1
        else:
            untested_ep += 1
            blind_spots.append(f"  ⚠ [{ep['id']}] {ep['method']} {ep['url']} — all pending")

    total_done = passed + failed + na_c + skipped_c
    cov_pct = (total_done / total_enabled * 100) if total_enabled else 0

    d = f"""
{'═' * 55}
        BOUNTYLENS COVERAGE DASHBOARD
{'═' * 55}

ENDPOINTS: {total_ep}
  ✅ Fully tested: {fully} | 🔄 Partial: {partially} | ❌ Untested: {untested_ep}

TESTS: {total_enabled} enabled (of {total_tests} total)
  ✅ Pass: {passed} | ❌ Fail: {failed} | ⬜ N/A: {na_c} | ⏭ Skip: {skipped_c} | 🔄 Pending: {pending}

COVERAGE: {cov_pct:.1f}%

{'─' * 55}
OWASP API TOP 10:
{'─' * 55}"""

    for cid, info in OWASP_CATEGORIES.items():
        c = owasp_cov[cid]
        pct = (c["done"] / c["total"] * 100) if c["total"] else 0
        bar = "█" * int(pct / 5) + "░" * (20 - int(pct / 5))
        fl = f" ({c['failed']} FAIL)" if c["failed"] else ""
        d += f"\n  {cid}: {info['name'][:35]:<35} [{bar}] {pct:.0f}% ({c['done']}/{c['total']}){fl}"

    if blind_spots:
        d += f"\n\nBLIND SPOTS ({len(blind_spots)}):\n" + "\n".join(blind_spots)
    else:
        d += "\n\n✅ No blind spots!"

    no_ctx = [ep for ep in endpoints_db.values() if not ep["business_context"]]
    if no_ctx:
        d += f"\n\nMISSING CONTEXT ({len(no_ctx)}):"
        for ep in no_ctx:
            d += f"\n  [{ep['id']}] {ep['method']} {ep['url']}"
    return d


@mcp.tool()
def export_report(format: str = "json") -> str:
    """Export report as json, word, or pdf."""
    if not endpoints_db:
        return "No data."

    ts = datetime.now().strftime("%Y%m%d_%H%M%S")

    if format == "json":
        fn = f"bountylens_report_{ts}.json"
        report = {
            "title": "BountyLens API Security Test Report",
            "generated_at": datetime.now().isoformat(),
            "owasp_categories": OWASP_CATEGORIES,
            "summary": _gen_summary(),
            "endpoints": list(endpoints_db.values()),
        }
        with open(fn, "w") as f:
            json.dump(report, f, indent=2, default=str)
        return f"Exported: {fn}"

    elif format == "word":
        try:
            from docx import Document
        except ImportError:
            return "pip install python-docx"
        fn = f"bountylens_report_{ts}.docx"
        doc = Document()
        doc.add_heading("BountyLens — API Security Report", 0)
        doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
        s = _gen_summary()
        doc.add_heading("Summary", 1)
        doc.add_paragraph(f"Endpoints: {s['total_endpoints']} | Tests: {s['total_enabled']} | Vulns: {s['failed']} | Coverage: {s['coverage_pct']:.1f}%")

        doc.add_heading("OWASP Coverage", 1)
        t = doc.add_table(rows=1, cols=4, style="Table Grid")
        for i, h in enumerate(["Category", "Tests", "Done", "Failures"]):
            t.rows[0].cells[i].text = h
        for cid, info in OWASP_CATEGORIES.items():
            c = s['owasp'].get(cid, {})
            row = t.add_row().cells
            row[0].text = f"{cid}: {info['name']}"
            row[1].text = str(c.get('total', 0))
            row[2].text = str(c.get('done', 0))
            row[3].text = str(c.get('failed', 0))

        for ep in endpoints_db.values():
            doc.add_heading(f"{ep['method']} {ep['url']}", 2)
            if ep["business_context"]:
                doc.add_paragraph(f"Context: {ep['business_context']}")
            doc.add_paragraph(f"Risk: {ep['risk_level'].upper()}")
            enabled = [tc for tc in ep["test_cases"] if tc.get("enabled")]
            if enabled:
                t = doc.add_table(rows=1, cols=5, style="Table Grid")
                for i, h in enumerate(["ID", "Test", "Severity", "Status", "Evidence"]):
                    t.rows[0].cells[i].text = h
                for tc in enabled:
                    row = t.add_row().cells
                    row[0].text = tc["id"]
                    row[1].text = tc["title"]
                    row[2].text = tc["severity"].upper()
                    row[3].text = tc["status"].upper()
                    row[4].text = tc.get("evidence", "") or "-"
        doc.save(fn)
        return f"Exported: {fn}"

    elif format == "pdf":
        try:
            from reportlab.lib.pagesizes import A4
            from reportlab.lib.styles import getSampleStyleSheet
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
            from reportlab.lib import colors
        except ImportError:
            return "pip install reportlab"
        fn = f"bountylens_report_{ts}.pdf"
        doc = SimpleDocTemplate(fn, pagesize=A4)
        styles = getSampleStyleSheet()
        els = [Paragraph("BountyLens — API Security Report", styles["Title"]), Spacer(1, 20)]
        s = _gen_summary()
        els.append(Paragraph(f"Endpoints: {s['total_endpoints']} | Tests: {s['total_enabled']} | Vulns: {s['failed']} | Coverage: {s['coverage_pct']:.1f}%", styles["Normal"]))
        els.append(Spacer(1, 20))
        for ep in endpoints_db.values():
            els.append(Paragraph(f"{ep['method']} {ep['url']}", styles["Heading2"]))
            if ep["business_context"]:
                els.append(Paragraph(f"Context: {ep['business_context']}", styles["Normal"]))
            enabled = [tc for tc in ep["test_cases"] if tc.get("enabled")]
            if enabled:
                data = [["ID", "Test", "Sev", "Status"]]
                for tc in enabled:
                    data.append([tc["id"], tc["title"][:45], tc["severity"][:4].upper(), tc["status"].upper()])
                t = Table(data, colWidths=[60, 230, 45, 55])
                t.setStyle(TableStyle([
                    ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#1a1a2e")),
                    ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                    ("FONTSIZE", (0, 0), (-1, -1), 7),
                    ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
                ]))
                els.append(t)
            els.append(Spacer(1, 14))
        doc.build(els)
        return f"Exported: {fn}"

    return f"Unknown format: {format}"


def _status_icon(s):
    return {"pass": "✅", "fail": "❌", "na": "⬜", "skipped": "⏭", "pending": "🔄"}.get(s, "❓")


def _gen_summary():
    tt = te = p = f = na = sk = pn = 0
    ow = {c: {"total": 0, "done": 0, "failed": 0} for c in OWASP_CATEGORIES}
    for ep in endpoints_db.values():
        for tc in ep["test_cases"]:
            tt += 1
            if not tc.get("enabled"): continue
            te += 1
            oc = tc.get("owasp", "")
            if oc in ow: ow[oc]["total"] += 1
            if tc["status"] == "pass": p += 1; ow.get(oc, {}) and oc in ow and (ow[oc].__setitem__("done", ow[oc]["done"] + 1))
            elif tc["status"] == "fail": f += 1; oc in ow and (ow[oc].__setitem__("done", ow[oc]["done"] + 1), ow[oc].__setitem__("failed", ow[oc]["failed"] + 1))
            elif tc["status"] == "na": na += 1; oc in ow and ow[oc].__setitem__("done", ow[oc]["done"] + 1)
            elif tc["status"] == "skipped": sk += 1; oc in ow and ow[oc].__setitem__("done", ow[oc]["done"] + 1)
            else: pn += 1
    dn = p + f + na + sk
    return {"total_endpoints": len(endpoints_db), "total_tests": tt, "total_enabled": te,
            "passed": p, "failed": f, "na": na, "skipped": sk, "pending": pn,
            "coverage_pct": (dn / te * 100) if te else 0, "owasp": ow}


if __name__ == "__main__":
    start_bridge_server()
    mcp.run()
